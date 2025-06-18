import streamlit as st
import pandas as pd
import plotly.express as px
from datetime import date
import os
import locale

# Réglage de la locale en français pour les noms de mois
try:
    locale.setlocale(locale.LC_TIME, "fr_FR.UTF-8")  # Linux/macOS
except locale.Error:
    try:
        locale.setlocale(locale.LC_TIME, "French_France")  # Windows, selon config
    except locale.Error:
        pass  # Locale non disponible, restera en anglais

st.set_page_config(page_title="Gestion des absences", layout="wide")

DATA_PATH = "absences.csv"

# Chargement ou initialisation du dataframe

if "df" not in st.session_state:
    if os.path.exists(DATA_PATH):
        st.session_state.df = pd.read_csv(DATA_PATH, parse_dates=["Date de l'absence"])
    else:
        st.session_state.df = pd.DataFrame(columns=[
            "Nom du patient",
            "Date de l'absence",
            "Statut de l'absence",
            "MAG",
            "Type de prise en charge",
            "Commentaire"
        ])

st.markdown(
    """
    <style>
    /* Fond principal bleu clair */
    .main {
        background-color: #e6f2ff;
        padding: 2rem 3rem;
        position: relative;
        min-height: 100vh;
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
    }

    /* Barre header plus foncée */
    header, footer, .css-18e3th9 {
        background-color: #cce0ff !important;
    }

    /* Motif ECG en filigrane */
    .main::before {
        content: "";
        position: absolute;
        top: 50px;
        left: 50%;
        transform: translateX(-50%);
        width: 200px;
        height: 200px;
        background-image: url('https://cdn.pixabay.com/photo/2016/10/13/12/42/heartbeat-1734588_960_720.png');
        background-repeat: no-repeat;
        background-size: contain;
        opacity: 0.1;
        pointer-events: none;
        z-index: 0;
    }

    /* Pour que le contenu soit au-dessus du filigrane */
    .block-container {
        position: relative;
        z-index: 1;
    }
    </style>
    """,
    unsafe_allow_html=True
)

st.markdown(
    """
    <div style="
        width: 100%;
        background-color: #f3e6ff;
        color: white;
        padding: 20px 0;
        text-align: center;
        font-size: 24px;
        font-weight: bold;
        ">
        Gestion des absences - Tableau de bord
    </div>
    """,
    unsafe_allow_html=True
)


st.title("📋 Tableau de bord de gestion des absences patients")
        
tabs = st.tabs([
    "📄 Gestion des absences",
    "📊 Statistiques",
    "📅 Statistiques annuelles",
    "🧾 Générateur de bilan"
])

# Onglet 1 : Gestion des absences
with tabs[0]:
    st.header("📄 Saisie d'une absence")
    with st.form("form_absence"):
        col1, col2, col3 = st.columns(3)
        with col1:
            nom = st.text_input("Nom du patient")
            date_abs = st.date_input("Date de l'absence", value=date.today())
        with col2:
            statut = st.selectbox("Statut de l'absence", ["Annulée", "Non venue", "Excusée", "Reportée"])
            mag = st.number_input("Manque à gagner (€)", min_value=0.0, step=1.0)
        with col3:
            type_pc = st.selectbox("Type de prise en charge", ["Consultation", "Suivi", "Urgence", "Bilan"])
            commentaire = st.text_area("Commentaire")

        submitted = st.form_submit_button("Ajouter l'absence")

        if submitted:
            new_row = pd.DataFrame([{
                "Nom du patient": nom,
                "Date de l'absence": pd.to_datetime(date_abs),
                "Statut de l'absence": statut,
                "MAG": mag,
                "Type de prise en charge": type_pc,
                "Commentaire": commentaire
            }])
            st.session_state.df = pd.concat([st.session_state.df, new_row], ignore_index=True)
            st.session_state.df.to_csv(DATA_PATH, index=False)
            st.success("Absence ajoutée avec succès !")

    st.subheader("🗂️ Liste des absences")
    edited_df = st.data_editor(
        st.session_state.df,
        num_rows="dynamic",
        use_container_width=True,
        key="data_editor"
    )
    st.session_state.df = edited_df
    st.session_state.df.to_csv(DATA_PATH, index=False)

# Onglet 2 : Statistiques
with tabs[1]:
    st.header("📊 Statistiques interactives")

    if st.session_state.df.empty:
        st.info("Aucune donnée disponible.")
    else:
        df = st.session_state.df.copy()
        df["Date de l'absence"] = pd.to_datetime(df["Date de l'absence"])
        df["Nom du patient"] = df["Nom du patient"].astype(str).str.strip()
        df["Statut de l'absence"] = df["Statut de l'absence"].astype(str).str.strip()

        st.subheader("📅 Filtrer par période")
        col1, col2 = st.columns(2)
        with col1:
            start_date = st.date_input("Date de début", value=df["Date de l'absence"].min().date())
        with col2:
            end_date = st.date_input("Date de fin", value=df["Date de l'absence"].max().date())

        mask = (df["Date de l'absence"] >= pd.to_datetime(start_date)) & \
               (df["Date de l'absence"] <= pd.to_datetime(end_date))
        filtered_df = df[mask]

        st.divider()
        st.subheader("📊 Absences par patient et par statut (barres empilées)")

        if not filtered_df.empty:
            absences_patient_statut = (
                filtered_df
                .groupby(["Nom du patient", "Statut de l'absence"])
                .size()
                .reset_index(name="Nombre d'absences")
            )

            fig1 = px.bar(
                absences_patient_statut,
                x="Nom du patient",
                y="Nombre d'absences",
                color="Statut de l'absence",
                text="Nombre d'absences",
                barmode="stack",
                template="plotly_white"
            )

            fig1.update_layout(
                xaxis_title="Patient",
                yaxis_title="Nombre d'absences",
                legend_title="Statut"
            )

            st.plotly_chart(fig1, use_container_width=True)
        else:
            st.warning("Aucune absence sur cette période.")

        st.divider()
        st.subheader("💸 Comparaison du manque à gagner (MAG)")

        patients = df["Nom du patient"].dropna().unique().tolist()
        selected_patient = st.selectbox("Choisir un patient", patients)

        if not filtered_df.empty:
            mag_total = filtered_df["MAG"].sum()
            mag_patient = filtered_df[filtered_df["Nom du patient"] == selected_patient]["MAG"].sum()

            mag_comparatif = pd.DataFrame({
                "Catégorie": [selected_patient, "Tous les patients"],
                "MAG (€)": [mag_patient, mag_total]
            })

            fig2 = px.bar(
                mag_comparatif,
                x="Catégorie",
                y="MAG (€)",
                color="Catégorie",
                text="MAG (€)",
                title="Manque à gagner du patient vs. total",
                template="plotly_white"
            )
            st.plotly_chart(fig2, use_container_width=True)
        else:
            st.warning("Aucun MAG à afficher sur cette période.")

# Onglet 3 : Statistiques annuelles
with tabs[2]:
    st.header("📅 Statistiques annuelles")

    if st.session_state.df.empty:
        st.info("Aucune donnée disponible.")
    else:
        df = st.session_state.df.copy()
        df["Date de l'absence"] = pd.to_datetime(df["Date de l'absence"])
        df["Année"] = df["Date de l'absence"].dt.year
        df["Mois"] = df["Date de l'absence"].dt.month

        # Graphique 1 : Manque à gagner total par année
        mag_par_annee = df.groupby("Année")["MAG"].sum().reset_index()

        fig_annee = px.bar(
            mag_par_annee,
            x="Année",
            y="MAG",
            color="Année",
            labels={"MAG": "Manque à gagner (€)", "Année": "Année"},
            title="Manque à gagner total par année",
            template="plotly_white"
        )

        fig_annee.update_xaxes(
            dtick=1,
            tickformat="d"
        )
        # Supprimer l'échelle de couleurs (colorbar)
        fig_annee.update_layout(coloraxis_showscale=False)

        st.plotly_chart(fig_annee, use_container_width=True)

        # Sélecteur d'année existante
        annees = mag_par_annee["Année"].sort_values().unique()
        annee_selectionnee = st.selectbox("Choisir une année pour détail mensuel", annees)

        # Filtrer pour l'année sélectionnée
        df_annee = df[df["Année"] == annee_selectionnee]

        # Tous les mois pour assurer affichage complet
        all_months = pd.DataFrame({"Mois": range(1, 13)})

        # Calcul MAG par mois
        mag_par_mois = df_annee.groupby("Mois")["MAG"].sum().reset_index()

        # Merge pour avoir tous les mois même sans données
        mag_par_mois = all_months.merge(mag_par_mois, on="Mois", how="left").fillna(0)

        # Convertir numéros en noms français minuscules
        mag_par_mois["Nom du mois"] = mag_par_mois["Mois"].apply(
            lambda x: pd.to_datetime(f"2023-{x:02d}-01").strftime('%B').lower()
        )

        ordre_mois = [
            "janvier", "février", "mars", "avril", "mai", "juin",
            "juillet", "août", "septembre", "octobre", "novembre", "décembre"
        ]

        couleurs_mois = {
            "janvier": "#636EFA",
            "février": "#EF553B",
            "mars": "#00CC96",
            "avril": "#AB63FA",
            "mai": "#FFA15A",
            "juin": "#19D3F3",
            "juillet": "#FF6692",
            "août": "#B6E880",
            "septembre": "#FF97FF",
            "octobre": "#FECB52",
            "novembre": "#636EFA",
            "décembre": "#EF553B",
        }

        fig_mois = px.bar(
            mag_par_mois,
            x="Nom du mois",
            y="MAG",
            labels={"MAG": "Manque à gagner (€)", "Nom du mois": "Mois"},
            title=f"Manque à gagner par mois en {annee_selectionnee}",
            template="plotly_white",
            color="Nom du mois",
            category_orders={"Nom du mois": ordre_mois},
            color_discrete_map=couleurs_mois
        )

        fig_mois.update_xaxes(tickangle=45)

        st.plotly_chart(fig_mois, use_container_width=True)

# Onglet 4 : Générateur de bilan
from docx import Document
from docx.shared import Pt
from io import BytesIO

# Exemple de structure de texte types

TEXTES_TYPES = {
    # Bilan type A
    ("Bilan type A", "épreuve 1.1", "faible"): "Résultat faible à l’épreuve 1.1 du bilan type A, indiquant des capacités limitées.",
    ("Bilan type A", "épreuve 1.1", "moyen"): "Résultat moyen à l’épreuve 1.1 du bilan type A, correspondant à une performance modérée.",
    ("Bilan type A", "épreuve 1.1", "bon"): "Bon résultat à l’épreuve 1.1 du bilan type A, reflétant des compétences solides.",
    
    ("Bilan type A", "épreuve 1.2", "faible"): "Performance faible à l’épreuve 1.2, ce qui suggère des difficultés notables.",
    ("Bilan type A", "épreuve 1.2", "moyen"): "Performance moyenne à l’épreuve 1.2, résultat dans la norme.",
    ("Bilan type A", "épreuve 1.2", "bon"): "Résultat satisfaisant à l’épreuve 1.2, montrant une bonne maîtrise.",
    
    ("Bilan type A", "épreuve 1.3", "faible"): "L’épreuve 1.3 révèle une performance faible, nécessitant une attention particulière.",
    ("Bilan type A", "épreuve 1.3", "moyen"): "L’épreuve 1.3 montre un niveau intermédiaire, sans difficulté majeure.",
    ("Bilan type A", "épreuve 1.3", "bon"): "L’épreuve 1.3 est réussie avec de bons résultats.",

    # Bilan type B
    ("Bilan type B", "épreuve 2.1", "faible"): "Résultat faible à l’épreuve 2.1 du bilan type B, en dessous des attentes.",
    ("Bilan type B", "épreuve 2.1", "moyen"): "Résultat moyen à l’épreuve 2.1, indiquant une performance acceptable.",
    ("Bilan type B", "épreuve 2.1", "bon"): "Très bonne performance à l’épreuve 2.1, dans les normes supérieures.",

    ("Bilan type B", "épreuve 2.2", "faible"): "Épreuve 2.2 difficile pour le patient, avec des résultats faibles.",
    ("Bilan type B", "épreuve 2.2", "moyen"): "Épreuve 2.2 réalisée avec des résultats moyens.",
    ("Bilan type B", "épreuve 2.2", "bon"): "Épreuve 2.2 réussie avec de bons résultats, conforme aux attentes.",

    ("Bilan type B", "épreuve 2.3", "faible"): "Des limites importantes sont observées à l’épreuve 2.3.",
    ("Bilan type B", "épreuve 2.3", "moyen"): "Résultat intermédiaire à l’épreuve 2.3.",
    ("Bilan type B", "épreuve 2.3", "bon"): "Bonne réussite de l’épreuve 2.3, sans difficulté apparente.",

    # Bilan type C
    ("Bilan type C", "épreuve 3.1", "faible"): "L’épreuve 3.1 a révélé un résultat faible, en deçà de la norme attendue.",
    ("Bilan type C", "épreuve 3.1", "moyen"): "Résultat moyen obtenu à l’épreuve 3.1.",
    ("Bilan type C", "épreuve 3.1", "bon"): "Excellente performance à l’épreuve 3.1, sans difficulté.",

    ("Bilan type C", "épreuve 3.2", "faible"): "Faiblesse marquée sur l’épreuve 3.2 du bilan C.",
    ("Bilan type C", "épreuve 3.2", "moyen"): "Résultat correct à l’épreuve 3.2.",
    ("Bilan type C", "épreuve 3.2", "bon"): "Très bon score à l’épreuve 3.2.",

    ("Bilan type C", "épreuve 3.3", "faible"): "Performance en difficulté à l’épreuve 3.3.",
    ("Bilan type C", "épreuve 3.3", "moyen"): "Performance dans la moyenne pour l’épreuve 3.3.",
    ("Bilan type C", "épreuve 3.3", "bon"): "Résultat élevé à l’épreuve 3.3, montrant de bonnes capacités.",
}


with tabs[3]:
    st.header("🧾 Générateur automatique de bilans médicaux")

    patient = st.selectbox("Nom du patient", ["Patient 1", "Patient 2", "Patient 3"])

    types_bilan = st.multiselect(
        "Types de bilan réalisés",
        ["Bilan type A", "Bilan type B", "Bilan type C"]
    )

    selections = {}

    for bilan in types_bilan:
        st.subheader(bilan)
        if bilan == "Bilan type A":
            epreuves = st.multiselect(f"Épreuves pour {bilan}", ["épreuve 1.1", "épreuve 1.2", "épreuve 1.3"], key=bilan)
        elif bilan == "Bilan type B":
            epreuves = st.multiselect(f"Épreuves pour {bilan}", ["épreuve 2.1", "épreuve 2.2", "épreuve 2.3"], key=bilan)
        elif bilan == "Bilan type C":
            epreuves = st.multiselect(f"Épreuves pour {bilan}", ["épreuve 3.1", "épreuve 3.2", "épreuve 3.3"], key=bilan)
        else:
            epreuves = []

        for epreuve in epreuves:
            result = st.selectbox(f"Résultat pour {epreuve}", ["faible", "moyen", "bon"], key=f"{bilan}-{epreuve}")
            selections[(bilan, epreuve)] = result

    if st.button("📄 Générer le bilan Word"):
        doc = Document()
        doc.add_heading(f"Bilan pour {patient}", 0)

        bilans_vus = set()

        for (bilan, epreuve), resultat in selections.items():
            if bilan not in bilans_vus:
                doc.add_heading(f"{bilan}", level=1)
                bilans_vus.add(bilan)

            doc.add_heading(f"{epreuve}", level=2)
            texte = TEXTES_TYPES.get((bilan, epreuve, resultat), "Texte non défini.")
            doc.add_paragraph(texte)

        # Sauvegarde dans un buffer mémoire
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            label="📥 Télécharger le bilan Word",
            data=buffer,
            file_name=f"bilan_{patient.replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
